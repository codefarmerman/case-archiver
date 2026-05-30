"""
auto_write.py — 缺件自动补写模块
目前支持：第 8 项代理词、第 11 项办案小结
需要设置环境变量 DEEPSEEK_API_KEY
"""
from __future__ import annotations

import os
import re
from pathlib import Path
from typing import List, Optional

from llm_client import LLMClient, extract_sample
from logger import get_logger

log = get_logger()

# 输出正文最小可接受长度（字符），低于此视为生成失败
MIN_CONTENT_CHARS = 100
MAX_WRITE_TOKENS = 8000

# 开头 AI 客套前言（仅匹配正文最开头）
_AI_PREAMBLE_PATTERNS = [
    r"^好的[，,。.！!]*\s*",
    r"^以下是.*?[:：]\s*",
    r"^这是.*?[:：]\s*",
]
# 结尾 AI 客套语（仅当作为最后一行整行出现时才剥离，避免误伤正文）
_AI_TRAILING_PHRASES = ("希望", "以上内容仅供参考", "如有需要")


def check_api_key() -> bool:
    return bool(os.environ.get("DEEPSEEK_API_KEY"))


def _clean_llm_output(text: str) -> str:
    """剥离 markdown 标记与 AI 客套前言，得到干净的文书正文。"""
    if not text:
        return ""

    # 去掉代码围栏 ```...```
    lines = [ln for ln in text.splitlines() if not ln.strip().startswith("```")]
    text = "\n".join(lines)

    # 去掉 markdown 标题井号、加粗/斜体、行首列表符号
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = text.replace("**", "").replace("__", "")
    text = re.sub(r"(?<!\*)\*(?!\*)", "", text)
    text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)

    # 去掉开头 AI 客套句
    stripped = text.strip()
    for pat in _AI_PREAMBLE_PATTERNS:
        stripped = re.sub(pat, "", stripped, flags=re.IGNORECASE).strip()

    # 去掉结尾客套语：仅当最后一非空行以客套词开头时剥离
    body_lines = stripped.splitlines()
    while body_lines:
        last = body_lines[-1].strip()
        if last and any(last.startswith(p) for p in _AI_TRAILING_PHRASES):
            body_lines.pop()
        else:
            break
    stripped = "\n".join(body_lines)

    return stripped.strip()


def auto_write_document(
    category_id: int,
    prompt_template: str,
    role: str,
    case_no: str,
    case_name: str,
    reference_files: List[Path],
    output_path: Path,
    llm: Optional[LLMClient] = None,
) -> Optional[Path]:
    """调用 Claude API 自动撰写文书。"""
    client = llm or LLMClient()
    if not client.ready:
        log.warning("跳过第 %d 项自动补写：%s", category_id, client.init_error or "LLM 未就绪")
        return None

    # 收集参考材料内容
    context_parts: List[str] = []
    for f in reference_files:
        text = extract_sample(f, n_chars=5000, max_pages=10)
        if text:
            context_parts.append(f"=== {f.name} ===\n{text}\n")

    if not context_parts:
        log.warning("第 %d 项无可用参考材料，跳过自动补写", category_id)
        return None

    context = "\n".join(context_parts)
    if len(context) > 15000:
        context = context[:15000] + "\n...(内容过长，已截断)"

    prompt = prompt_template.format(role=role)
    full_prompt = (
        f"{prompt}\n\n"
        f"案号：{case_no}\n"
        f"案由/当事人：{case_name}\n"
        f"我方诉讼地位：{role}\n\n"
        f"以下是案件相关材料：\n{context}\n\n"
        "请直接输出文书正文，不要输出标题外的额外说明。"
    )

    log.info("正在调用 DeepSeek 撰写第 %d 项 ...", category_id)
    content = client.write_document(
        system="你是一名资深中国律师，擅长撰写标准诉讼文书。直接输出正文，不要使用 markdown 标记，不要客套开场白。",
        prompt=full_prompt,
        max_tokens=MAX_WRITE_TOKENS,
    )
    if not content:
        log.error("第 %d 项 API 调用失败", category_id)
        return None

    content = _clean_llm_output(content)
    if len(content) < MIN_CONTENT_CHARS:
        log.warning("第 %d 项生成内容过短（%d 字），视为失败，不写文件", category_id, len(content))
        return None

    try:
        _write_docx(content, output_path, case_no, category_id)
        log.info("已生成：%s（%d 字）", output_path.name, len(content))
        return output_path
    except Exception as e:
        log.error("写入 docx 失败：%s", e)
        return None


def _write_docx(content: str, output_path: Path, case_no: str, category_id: int):
    """将文本内容写入 docx。"""
    try:
        from docx import Document as DocxDocument
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        txt_path = output_path.with_suffix(".txt")
        txt_path.write_text(content, encoding="utf-8")
        log.warning("python-docx 未安装，已输出为 txt：%s", txt_path.name)
        return

    doc = DocxDocument()
    titles = {8: "代 理 词", 11: "办 案 小 结"}
    title = doc.add_heading(titles.get(category_id, ""), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"案号：{case_no}")
    run.font.size = Pt(12)

    for para_text in content.split("\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph(para_text)
        for r in p.runs:
            r.font.size = Pt(12)
            r.font.name = "宋体"

    doc.save(str(output_path))
