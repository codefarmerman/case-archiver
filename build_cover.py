"""
build_cover.py — 生成卷内目录Word v2.0
按13项标准结构生成目录表，标注我方/对方，空项也列出。
"""
from __future__ import annotations

from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

CATEGORY_DISPLAY = {
    1:  "收案审查表",
    2:  "收费凭证",
    3:  "法律事务委托合同、授权委托书",
    4:  "起诉书、申请书、上诉书、答辩书、反诉状、再审申请书、抗诉书",
    5:  "阅卷笔录",
    6:  "证据材料",
    7:  "出庭通知书",
    8:  "代 理 词",
    9:  "庭审笔录",
    10: "判决书、裁定书、裁决书、调解书、上诉书",
    11: "办案小结",
    12: "服务质量意见反馈表",
    13: "其    他",
}


def _set_cell_font(cell, text: str, size: int = 10, bold: bool = False):
    """设置单元格字体。"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def build_cover(
    output_path: Path,
    case_no: str,
    case_name: str,
    classified_files: List[dict],
):
    """生成卷内目录Word。

    classified_files: [{"category_id":int, "filename":str, "side":str, "page_start":str}]
    """
    doc = Document()

    # ---- 标题 ----
    title = doc.add_heading("", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("民事、刑附民、仲裁、行政案卷卷内目录")
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(16)

    # ---- 案件信息 ----
    info = doc.add_paragraph()
    run1 = info.add_run(f"案    号：{case_no}")
    run1.font.size = Pt(12)
    run1.font.name = "宋体"
    info.add_run("\n")
    run2 = info.add_run(f"案由/当事人：{case_name}")
    run2.font.size = Pt(12)
    run2.font.name = "宋体"

    # ---- 主表 ----
    # 4列：序号 | 案卷材料 | 实际文件 | 页码
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # 表头
    headers = ["序号", "案卷材料", "实际文件", "页码"]
    for i, h in enumerate(headers):
        _set_cell_font(table.rows[0].cells[i], h, size=10, bold=True)

    # 按13项顺序填表
    files_by_cat = {}
    for f in classified_files:
        files_by_cat.setdefault(f["category_id"], []).append(f)

    for cat_id in range(1, 14):
        files = files_by_cat.get(cat_id, [])
        display_name = CATEGORY_DISPLAY[cat_id]

        if not files:
            row = table.add_row().cells
            _set_cell_font(row[0], str(cat_id))
            _set_cell_font(row[1], display_name)
            _set_cell_font(row[2], "（无）")
            _set_cell_font(row[3], "")
        else:
            for idx, f in enumerate(files):
                row = table.add_row().cells
                seq = f"{cat_id}" + (f"-{idx+1}" if len(files) > 1 else "")
                _set_cell_font(row[0], seq)
                _set_cell_font(row[1], display_name if idx == 0 else "")

                # 文件名（第4项加我方/对方标注）
                fname = f["filename"]
                side = f.get("side", "")
                display = f"[{side}] {fname}" if side and side != "待确认" else fname
                _set_cell_font(row[2], display)

                _set_cell_font(row[3], str(f.get("page_start", "")))

    # ---- 备考 ----
    row = table.add_row().cells
    _set_cell_font(row[0], "备考")
    # 合并后3列
    row[1].merge(row[2]).merge(row[3])

    doc.save(output_path)
    return output_path
