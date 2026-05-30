"""卷内目录生成测试。"""
from docx import Document

from build_cover import build_cover


def test_build_cover_creates_file(tmp_path):
    out = tmp_path / "卷内目录.docx"
    files = [
        {"category_id": 4, "filename": "起诉状.pdf", "side": "我方", "page_start": ""},
        {"category_id": 4, "filename": "答辩状.pdf", "side": "对方", "page_start": ""},
    ]
    result = build_cover(out, "(2024)案号001", "张三诉李四", files)
    assert result.exists()


def test_build_cover_lists_all_13_categories(tmp_path):
    out = tmp_path / "卷内目录.docx"
    build_cover(out, "案号", "案由", [])
    doc = Document(str(out))
    # 应包含所有 13 项（空项显示"（无）"）
    table = doc.tables[0]
    text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    assert "收案审查表" in text
    assert "办案小结" in text
    assert "（无）" in text  # 空项标记


def test_build_cover_side_annotation(tmp_path):
    out = tmp_path / "卷内目录.docx"
    files = [{"category_id": 4, "filename": "起诉状.pdf", "side": "我方", "page_start": ""}]
    build_cover(out, "案号", "案由", files)
    doc = Document(str(out))
    table = doc.tables[0]
    text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    assert "[我方]" in text


def test_build_cover_daiqueren_no_annotation(tmp_path):
    # 待确认不应加 [side] 标注
    out = tmp_path / "卷内目录.docx"
    files = [{"category_id": 4, "filename": "某文书.pdf", "side": "待确认", "page_start": ""}]
    build_cover(out, "案号", "案由", files)
    doc = Document(str(out))
    table = doc.tables[0]
    text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    assert "[待确认]" not in text
